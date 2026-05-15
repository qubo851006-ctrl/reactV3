from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = PROJECT_ROOT / "review_outputs"
OUTPUT_PATH = OUTPUT_DIR / "reactV2_code_review_report.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.22
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Title", 22, "1F4E79"),
        ("Heading 1", 15, "1F4E79"),
        ("Heading 2", 12, "2F5597"),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def add_meta_table(doc: Document) -> None:
    rows = [
        ("项目路径", r"D:\claude\reactV2"),
        ("报告日期", "2026-04-28"),
        ("审查类型", "代码 Review：安全风险、数据完整性、质量门禁"),
        ("审查结论", "发现 4 项 P1 高风险问题、2 项 P2 中风险问题；前端构建通过，lint 未通过，后端语法编译通过。"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(1.3)
    table.columns[1].width = Inches(5.6)
    for i, (label, value) in enumerate(rows):
        set_cell_text(table.cell(i, 0), label, bold=True, color="1F4E79")
        set_cell_shading(table.cell(i, 0), "D9EAF7")
        set_cell_text(table.cell(i, 1), value)


def add_summary_table(doc: Document) -> None:
    findings = [
        ("P1", "session_id 路径穿越", "backend/routers/chat.py:157", "可读写/删除用户历史或覆盖 data 下 JSON 文件", "白名单校验 session_id，并校验 resolve 后路径边界"),
        ("P1", "Ledger 日志 XSS", "frontend/src/components/LedgerFlow.tsx:93", "上传文件名进入 dangerouslySetInnerHTML，可在登录态执行脚本", "移除 innerHTML，改为 React 节点或严格 HTML 转义"),
        ("P1", "上传文件名路径逃逸", "backend/routers/training.py:41; backend/ledger_helpers.py:325", "恶意文件名可能写出临时目录或归档目录", "统一 basename + 文件名净化 + 最终路径边界校验"),
        ("P1", "4 位短码可枚举", "backend/routers/auth.py:15; backend/routers/auth.py:43", "公开用户列表且无失败限速，短码空间仅约 9000", "增加限流、失败锁定、审计，提升短码强度或接入正式认证"),
        ("P2", "台账并发写入风险", "backend/ledger_helpers.py:239; backend/utils/excel_writer.py:50", "多用户同时写 JSON/Excel 可能丢数据或损坏文件", "增加文件锁、原子写，或迁移到数据库事务"),
        ("P2", "TLS 校验关闭", "backend/llm_client.py:10; backend/utils/mcp_client.py:110", "外部 LLM/MCP 请求存在中间人风险", "配置企业可信 CA，移除 verify=False"),
    ]

    doc.add_heading("问题总览", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["等级", "问题", "位置", "影响", "整改建议"]
    widths = [0.65, 1.45, 2.0, 2.2, 2.2]
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_text(cell, header, bold=True, color="FFFFFF")
        set_cell_shading(cell, "1F4E79")
        table.columns[idx].width = Inches(widths[idx])

    for severity, title, location, impact, fix in findings:
        cells = table.add_row().cells
        values = [severity, title, location, impact, fix]
        for idx, value in enumerate(values):
            set_cell_text(cells[idx], value, bold=(idx == 0))
            if severity == "P1":
                set_cell_shading(cells[idx], "FCE4D6")
            else:
                set_cell_shading(cells[idx], "FFF2CC")


def add_detail(doc: Document, title: str, risk: str, evidence: str, suggestion: str) -> None:
    doc.add_heading(title, level=2)
    for label, text in [("风险说明", risk), ("代码证据", evidence), ("整改建议", suggestion)]:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}：")
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("1F4E79")
        p.add_run(text)


def build() -> Path:
    OUTPUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    set_doc_defaults(doc)

    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("reactV2 项目代码 Review 报告")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("安全风险、数据完整性与工程质量专项审查")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("666666")

    add_meta_table(doc)

    doc.add_heading("执行摘要", level=1)
    doc.add_paragraph(
        "本次审查覆盖后端 FastAPI/SQLite 文件型台账、前端 React/Vite 交互层、登录短码机制、文件上传下载链路以及现有质量门禁。"
        "整体判断：项目已具备基本业务闭环，但当前存在多处可由已登录用户或上传入口触发的高风险问题，尤其是路径穿越、XSS、上传文件名落盘和短码枚举风险，应优先进入修复排期。"
    )
    doc.add_paragraph(
        "建议优先级：先修复 4 项 P1 问题，再补齐并发写保护、TLS 校验和 lint 质量门禁。修复后应增加针对路径边界、恶意文件名、XSS payload、登录失败限速的回归测试。"
    )

    doc.add_page_break()
    add_summary_table(doc)

    doc.add_heading("详细发现", level=1)
    add_detail(
        doc,
        "1. [P1] session_id 可路径穿越并覆盖数据文件",
        "session_id 被直接拼接为聊天历史文件名。接口允许外部传入 session_id 时，攻击者可以构造相对路径，读写或删除当前用户目录之外的 JSON 文件。",
        "backend/routers/chat.py:157 直接返回 _user_dir(user_id) / f\"{session_id}.json\"；load_history、save_history、delete_session、clear_history 后续均复用该路径。",
        "只允许服务端创建的 session id；对所有入口使用正则白名单；Path.resolve 后确认最终路径位于 _user_dir(user_id) 内；对不存在于 sessions.json 的 session 拒绝操作。",
    )
    add_detail(
        doc,
        "2. [P1] 文件名进入 innerHTML 导致 XSS",
        "LedgerFlow 使用 dangerouslySetInnerHTML 渲染日志，日志中包含后端回传的原始上传文件名。恶意文件名可注入 HTML/JS，在登录态页面内执行敏感操作。",
        "backend/routers/ledger.py:49 把 fd['name'] 写入 SSE 日志；frontend/src/components/LedgerFlow.tsx:93 使用 dangerouslySetInnerHTML 渲染。",
        "删除 dangerouslySetInnerHTML；把日志拆成结构化事件，在 React 中按文本节点渲染；若必须支持 Markdown 样式，先做 HTML escape，再通过受信任渲染器处理有限语法。",
    )
    add_detail(
        doc,
        "3. [P1] 上传文件名可写出预期目录",
        "UploadFile.filename 被直接用于 os.path.join 或归档文件名，未做 basename 和路径边界校验。恶意文件名可导致写入位置逃逸。",
        "backend/routers/training.py:41-42 将 notice_pdf.filename/signin_img.filename 拼到 tmpdir；backend/ledger_helpers.py:325-333 使用 fd['name'] 生成归档路径。",
        "统一实现 safe_upload_filename；仅保留文件名主体和允许扩展名；写入前 resolve 并校验路径位于目标目录内；拒绝绝对路径、..、控制字符和过长文件名。",
    )
    add_detail(
        doc,
        "4. [P1] 4 位短码无防爆破保护",
        "登录页公开可选用户列表，短码仅 4 位且 bind-device 无限速、失败计数、锁定或审计。短码可被在线枚举。",
        "backend/routers/auth.py:15 users-lite 公开 active 用户；backend/routers/auth.py:43 校验短码；backend/routers/admin_users.py:59/107 生成 1000-9999 短码。",
        "增加按 IP、用户、设备维度的限流和失败锁定；失败写审计；短码提升到 6 位以上并设置有效期；条件允许时改为企业统一身份认证。",
    )
    add_detail(
        doc,
        "5. [P2] 台账文件并发写入可能丢数据",
        "案件台账 JSON、培训 Excel、授权台账和合并台账均采用文件读改写。多请求并发时，后写可能覆盖先写，Excel 文件也可能被同时写坏。",
        "backend/ledger_helpers.py:239/247 读写 cases.json；backend/utils/excel_writer.py:50/67 读写培训统计表；backend/routers/ledger.py:130-146 读改写案件台账。",
        "为每类台账增加文件锁；写入采用临时文件 + os.replace 原子替换；业务上需要多人协作时，将台账主数据迁到数据库并使用事务。",
    )
    add_detail(
        doc,
        "6. [P2] 外部调用关闭 TLS 校验",
        "多个 httpx/OpenAI 客户端使用 verify=False，法务文书、企业查询和 API key 流量存在被中间人篡改或窃取的风险。",
        "backend/llm_client.py:10、backend/utils/mcp_client.py:110、backend/routers/chat.py:360 均关闭证书校验。",
        "配置内网代理或企业根证书路径；将 CA 路径放入环境变量；禁止默认 verify=False，只在受控调试场景通过显式配置开启。",
    )

    doc.add_heading("验证记录", level=1)
    verifications = [
        ("前端构建", "npm.cmd run build", "通过；Vite 产物生成成功，但主 JS chunk 约 752.71 kB，后续建议 code splitting。"),
        ("前端 lint", "npm.cmd run lint", "未通过；共 11 个 @typescript-eslint/no-explicit-any 错误，集中在 api.ts 和多个 Flow 组件。"),
        ("后端语法编译", "python -m compileall -q backend", "通过；使用 Codex bundled Python 运行。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for idx, header in enumerate(["检查项", "命令", "结果"]):
        set_cell_text(table.cell(0, idx), header, bold=True, color="FFFFFF")
        set_cell_shading(table.cell(0, idx), "1F4E79")
    for item, command, result in verifications:
        row = table.add_row().cells
        set_cell_text(row[0], item, bold=True)
        set_cell_text(row[1], command)
        set_cell_text(row[2], result)

    doc.add_heading("建议整改顺序", level=1)
    for item in [
        "第一优先级：修复 session_id 路径穿越、Ledger 日志 XSS、上传文件名路径逃逸、短码防爆破。",
        "第二优先级：补齐台账并发写保护与 TLS 证书校验，避免数据损坏和外部调用链路风险。",
        "第三优先级：修复前端 lint 中的 any 类型问题，将 lint 纳入提交或发布前检查。",
        "第四优先级：为路径、上传、XSS、登录失败限速和并发写入增加回归测试。",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("reactV2 代码 Review 报告 | 2026-04-28").font.size = Pt(9)

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build())
