# -*- coding: utf-8 -*-
"""Utilities for the authorization request workflow.

The current authorization flow is intentionally template driven. Attachment 1
provides the project/source basis, Attachment 2 provides the authorization
fields, and a small set of user-entered fields completes the final request and
ledger row.
"""

from __future__ import annotations

import base64
import io
import logging
import os
import re
import shutil
import subprocess
import tempfile
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from typing import Any

import pdfplumber

logger = logging.getLogger(__name__)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from file_store import atomic_save_workbook, file_lock


AUTH_LEDGER_HEADERS = [
    "序号",
    "编号",
    "经办人",
    "授权人",
    "代理人",
    "印章",
    "份数",
    "授权起止日期",
    "授权内容概要",
    "办理时间",
    "代理人签字版是否发送回来",
    "文号",
    "责任者",
    "题目",
    "归档日期",
    "页数",
]


def _clean_text(text: str) -> str:
    text = text.replace("\x07", "\n").replace("\r", "\n")
    text = re.sub(r"0000060905\s+\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}:\d{2}", "", text)
    text = re.sub(r"\n(?:\s*[0-9][:：]?\s*){3,}\n", "\n", text)
    text = re.sub(r"\n(?:\s*[-—]\s*){3,}\n", "\n", text)
    text = re.sub(r"[ \t\u3000]+", " ", text)
    text = re.sub(r"\n\s+", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _strip_word_field_codes(text: str) -> str:
    text = re.sub(r"\x13[^\x14\x15]{0,1000}\x14", "", text)
    return text.replace("\x15", "")


def _compact(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def _normalize_doc_no(value: str | None) -> str | None:
    if not value:
        return None
    value = re.sub(r"\s+", "", value)
    value = value.replace("）", "〕").replace("（", "〔")
    m = re.search(r"(中航集团规划发〔\d{4}〕\d+号)", value)
    if m:
        return m.group(1)
    m = re.search(r"([\u4e00-\u9fa5A-Za-z]{2,12}发〔\d{4}〕\d+号)$", value)
    if m:
        return m.group(1)
    return value


def _normalize_project_name(value: str | None) -> str | None:
    if not value:
        return None
    value = re.sub(r"\s+", "", value)
    value = re.sub(r"^关于", "", value)
    value = re.sub(r"(需求调整及明确相关工作的通知|授权委托书|的请示|请示)$", "", value)
    if len(value) > 80:
        return None
    return value or None


def _safe_extracted_text(value: str | None, max_len: int, fallback: str | None = None) -> str | None:
    value = _dedupe_repeated(value)
    if not value:
        return fallback
    value = re.sub(r"\s+", "", value)
    return value if len(value) <= max_len else fallback


def _extract_project_from_scope(scope: str | None) -> str | None:
    if not scope:
        return None
    compact = _compact(scope)
    m = re.search(r"为开展(.+?)(?:项目|工程)(?:的)?相关工作", compact)
    if m:
        return _normalize_project_name(m.group(1) + "项目")
    return None


def _dedupe_repeated(value: str | None) -> str | None:
    if not value:
        return value
    value = value.strip()
    for _ in range(4):
        if len(value) % 2 == 0:
            half = len(value) // 2
            if value[:half] == value[half:]:
                value = value[:half]
                continue
        break
    sentences = [s for s in re.split(r"(?<=。)", value) if s]
    if len(sentences) > 1:
        compacted: list[str] = []
        for sentence in sentences:
            if sentence not in compacted:
                compacted.append(sentence)
        value = "".join(compacted)
    return value.strip()


def _normalize_authorization_term(value: str | None) -> str | None:
    value = _dedupe_repeated(value)
    if not value:
        return value
    value = re.sub(r"^\s*授权期限\s*[:：]?\s*", "", value)
    return value.strip()


def _short_name(full_name: str | None, fallback: str = "") -> str:
    if not full_name:
        return fallback
    if "中国航空集团建设开发有限公司" in full_name or "中航集团建设开发有限公司" in full_name:
        return "建开公司"
    if "中国国际航空股份有限公司" in full_name:
        return "国航股份"
    if full_name.endswith("有限公司") and len(full_name) <= 12:
        return full_name
    return full_name


def _extract_pdf_text(pdf_bytes: bytes) -> str:
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc)
        if text.strip():
            return _clean_text(text)
    except Exception:
        logger.debug("PyMuPDF 提取 PDF 文本失败,回退到 pdfplumber", exc_info=True)

    text = ""
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text:
                text += page_text + "\n"
    return _clean_text(text)


def _extract_docx_text(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return _clean_text("\n".join(parts))


def _extract_doc_text_from_binary(doc_bytes: bytes) -> str:
    text = _strip_word_field_codes(doc_bytes.decode("utf-16le", errors="ignore"))
    markers = ["授权 委 托 书", "授权委托书", "企业名称", "委托单位"]
    starts = [text.find(marker) for marker in markers if text.find(marker) >= 0]
    if starts:
        text = text[min(starts):]
    text = re.split(r"\x00{8,}", text, maxsplit=1)[0]
    text = _clean_text(text)
    compact = _compact(text)
    if "企业名称" not in compact or "授权编号" not in compact:
        raise RuntimeError(".doc 二进制文本兜底未提取到授权委托书正文")
    return text


def _resolve_soffice_binary() -> str | None:
    configured = os.getenv("LIBREOFFICE_BIN", "").strip()
    if configured:
        return configured
    return shutil.which("soffice") or shutil.which("libreoffice")


def _extract_doc_text_with_soffice(doc_bytes: bytes) -> str:
    soffice = _resolve_soffice_binary()
    if not soffice:
        raise RuntimeError(
            "服务器未配置 LibreOffice，无法解析 .doc 文件；"
            "请联系管理员配置 LIBREOFFICE_BIN，或上传 PDF/DOCX 格式附件2。"
        )
    with tempfile.TemporaryDirectory() as tmpdir:
        doc_path = os.path.join(tmpdir, "attachment2.doc")
        with open(doc_path, "wb") as f:
            f.write(doc_bytes)
        subprocess.run(
            [soffice, "--headless", "--convert-to", "docx", "--outdir", tmpdir, doc_path],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=60,
        )
        docx_path = os.path.join(tmpdir, "attachment2.docx")
        with open(docx_path, "rb") as f:
            return _extract_docx_text(f.read())


def extract_attachment_text(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf_text(file_bytes)
    if ext == ".docx":
        return _extract_docx_text(file_bytes)
    if ext == ".doc":
        try:
            return _extract_doc_text_from_binary(file_bytes)
        except Exception:
            logger.debug(".doc 二进制文本提取失败,回退到 LibreOffice 转换", exc_info=True)
        return _extract_doc_text_with_soffice(file_bytes)
    raise ValueError("不支持的附件格式")


def extract_attachment1_info(pdf_bytes: bytes, *, ocr_text: str | None = None) -> dict[str, Any]:
    text = _extract_pdf_text(pdf_bytes) or _clean_text(ocr_text or "")
    compact = _compact(text)

    title = _extract_attachment1_title(text)
    if not title:
        m = re.search(r"(关于.{3,80}?通知)", compact)
        if m:
            title = m.group(1)

    doc_no = None
    for pattern in [
        r"([\u4e00-\u9fa5A-Za-z]{2,20}规划发〔\d{4}〕\d+号)",
        r"([\u4e00-\u9fa5A-Za-z]{2,20}发〔\d{4}〕\d+号)",
    ]:
        matches = re.findall(pattern, compact)
        if matches:
            doc_no = _normalize_doc_no(min(matches, key=len))
            break

    project_name = _normalize_project_name(title)
    undertaking_unit = None
    undertaking_short = None
    if "中国航空集团建设开发有限公司" in compact:
        undertaking_unit = "中国航空集团建设开发有限公司"
        undertaking_short = "建开公司"
    else:
        m = re.search(r"同时请(.{4,40}?有限公司).*?开展前期工作", compact)
        if m:
            undertaking_unit = m.group(1)
            undertaking_short = _short_name(undertaking_unit)

    return {
        "raw_text": text,
        "title": title,
        "document_no": doc_no,
        "project_name": project_name,
        "undertaking_unit": undertaking_unit,
        "undertaking_short": undertaking_short,
    }


def _extract_attachment1_title(text: str) -> str | None:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for idx, line in enumerate(lines):
        normalized = re.sub(r"\s+", "", line)
        normalized = re.sub(r"^[0-9:：\\-—]+", "", normalized)
        if "关于" not in normalized or "《关于" in normalized:
            continue
        if len(normalized) > 60 and "通知" not in normalized:
            continue
        parts = [normalized[normalized.index("关于"):]]
        for next_line in lines[idx + 1: idx + 5]:
            next_norm = re.sub(r"\s+", "", next_line)
            next_norm = re.sub(r"^[0-9:：\\-—]+", "", next_norm)
            if not next_norm or re.fullmatch(r"[0-9:：\\-—]+", next_norm):
                continue
            parts.append(next_norm)
            if "通知" in next_norm:
                title = "".join(parts)
                title = re.sub(r"[0-9:：\\-—]{3,}", "", title)
                m = re.search(r"(关于.{3,80}?通知)", title)
                return m.group(1) if m else title
            if len("".join(parts)) > 80:
                break
    return None


def _field_between(compact: str, start: str, end: str) -> str | None:
    m = re.search(re.escape(start) + r"(.+?)" + re.escape(end), compact)
    return m.group(1).strip() if m else None


def _trim_repeated_label(value: str | None, *labels: str) -> str | None:
    if not value:
        return None
    for label in labels:
        value = value.replace(label, "")
    return value.strip() or None


def extract_attachment2_info(file_bytes: bytes, filename: str, *, ocr_text: str | None = None) -> dict[str, Any]:
    ext = Path(filename).suffix.lower()
    text = _clean_text(ocr_text or "") if ext == ".pdf" and ocr_text else extract_attachment_text(file_bytes, filename)
    compact = _compact(text)

    auth_no = None
    m = re.search(r"授权编号[:：]\s*([^\n\r]+)", text)
    if m:
        auth_no = re.sub(r"\s+", "", m.group(1)).strip()
    if not auth_no:
        m = re.search(r"(建开转托字[（(]\d{4}[）)]\s*\d+\s*号)", compact)
        if m:
            auth_no = _normalize_doc_no(m.group(1))

    principal = _dedupe_repeated(_field_between(compact, "企业名称", "注册地址"))
    legal_rep = _dedupe_repeated(_field_between(compact, "法定代表人", "职务"))
    trustee_name = _dedupe_repeated(_field_between(compact, "姓名", "电话"))
    trustee_phone = _dedupe_repeated(_field_between(compact, "电话", "工作单位"))
    trustee_work_unit = _dedupe_repeated(_field_between(compact, "工作单位", "职务"))
    trustee_position = None
    m = re.search(r"工作单位.+?职务(.+?)委托事项及权限", compact)
    if m:
        trustee_position = _dedupe_repeated(m.group(1).strip())
    scope = _dedupe_repeated(_field_between(compact, "委托事项及权限", "授权期限"))
    term = _normalize_authorization_term(_field_between(compact, "授权期限", "委托单位盖章"))
    remark = _dedupe_repeated(_field_between(compact, "备注说明", "授权编号"))

    scope = _trim_repeated_label(scope, "为开展")
    if scope and not scope.startswith("为开展"):
        scope = "为开展" + scope
    permission_detail = None
    if scope:
        m = re.search(r"其权限为[:：](.+?。?$)", scope)
        permission_detail = m.group(1).strip("。 ") if m else scope

    permission_type = "合同签署权限" if permission_detail and "合同" in permission_detail and "签署" in permission_detail else "授权权限"

    return {
        "raw_text": text,
        "authorization_no": _dedupe_repeated(auth_no),
        "principal_unit": principal,
        "principal_short": _short_name(principal, "委托单位"),
        "legal_representative": legal_rep,
        "trustee_name": trustee_name,
        "trustee_phone": trustee_phone,
        "trustee_work_unit": trustee_work_unit,
        "trustee_position": trustee_position,
        "authorization_scope": scope,
        "permission_detail": permission_detail,
        "permission_type": permission_type,
        "authorization_term": term,
        "copies_from_attachment": remark,
    }


def build_auth_request_text(extracted: dict[str, Any], user_inputs: dict[str, Any]) -> str:
    a1 = extracted.get("attachment1") or {}
    a2 = extracted.get("attachment2") or {}

    project_name = _normalize_project_name(a1.get("project_name")) or _extract_project_from_scope(a2.get("authorization_scope")) or "项目"
    source_title = _safe_extracted_text(a1.get("title"), 100)
    if not source_title:
        source_title = f"关于{project_name}需求调整及明确相关工作的通知" if project_name != "项目" else "相关文件"
    source_no = _normalize_doc_no(a1.get("document_no")) or "文件编号待补充"
    undertaking_unit = _safe_extracted_text(a1.get("undertaking_unit"), 60) or a2.get("principal_unit") or "承办单位"
    undertaking_short = _safe_extracted_text(a1.get("undertaking_short"), 20) or _short_name(undertaking_unit, "承办单位")

    trustee_work_unit = a2.get("trustee_work_unit") or ""
    if trustee_work_unit.startswith(a2.get("principal_unit") or ""):
        trustee_work_unit = trustee_work_unit.replace(a2.get("principal_unit") or "", _short_name(a2.get("principal_unit")), 1)
    trustee_position = _dedupe_repeated(a2.get("trustee_position")) or ""
    trustee_name = _dedupe_repeated(a2.get("trustee_name")) or "受托人"
    permission_type = a2.get("permission_type") or "授权权限"
    permission_detail = (_dedupe_repeated(a2.get("permission_detail")) or "授权内容待补充").rstrip("。")
    principal_short = _dedupe_repeated(a2.get("principal_short")) or _short_name(a2.get("principal_unit"), "委托单位")

    auth_mode = user_inputs.get("auth_mode")
    transfer_subject = (user_inputs.get("transfer_subject") or "").strip()
    if auth_mode == "transfer":
        procedure = f"，并履行{transfer_subject}转授权程序"
    else:
        procedure = ""

    copies = (user_inputs.get("copies") or "").strip()
    seal = (user_inputs.get("seal") or "").strip()

    return "\n".join(
        [
            f"依据《{source_title}》（{source_no}，详见附件1），{project_name}由{undertaking_unit}（以下简称{undertaking_short}）开展前期工作。",
            f"现拟将{project_name}的{permission_type}授予{trustee_work_unit}{trustee_position}{trustee_name}同志，需由{principal_short}出具授权委托书{procedure}。{trustee_name}具体授权内容:{permission_detail}。",
            f"授权委托期限:自审批通过之日起至{project_name}结束止。",
            f"授权办理份数：{copies}。",
            f"办理授权时需使用{seal}。",
            "妥否，请批示。",
            f"附件：1.{source_title}",
            "2.授权委托书",
        ]
    )


def _set_run_font(run, cn_font: str = "仿宋_GB2312", size_pt: float = 16) -> None:
    run.font.size = Pt(size_pt)
    run.font.name = cn_font
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), cn_font)
    r_fonts.set(qn("w:hAnsi"), cn_font)
    r_fonts.set(qn("w:eastAsia"), cn_font)
    r_fonts.set(qn("w:cs"), cn_font)


def _format_request_paragraph(paragraph, *, first_line: bool = True, before_pt: float = 0) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(31.2)
    fmt.space_before = Pt(before_pt)
    fmt.space_after = Pt(0)
    if first_line:
        fmt.first_line_indent = Pt(32)


def save_auth_request_docx(content: str, output_path: str) -> None:
    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    styles = doc.styles
    styles["Normal"].font.name = "仿宋_GB2312"
    styles["Normal"].font.size = Pt(16)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "仿宋_GB2312")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "仿宋_GB2312")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:cs"), "仿宋_GB2312")

    lines = [line.strip() for line in content.splitlines() if line.strip()]
    for idx, line in enumerate(lines):
        is_attachment = line.startswith("附件：")
        is_attachment_cont = re.match(r"^\d+\.", line) is not None
        p = doc.add_paragraph()
        if is_attachment:
            _format_request_paragraph(p, first_line=False, before_pt=19)
        elif is_attachment_cont:
            _format_request_paragraph(p, first_line=False)
            p.paragraph_format.left_indent = Pt(48)
        else:
            _format_request_paragraph(p, first_line=True)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        _set_run_font(run)

    doc.save(output_path)


def docx_to_base64(content: str, filename_prefix: str = "授权请示") -> tuple[str, str]:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False, prefix=f"{filename_prefix}_") as tmp:
        path = tmp.name
    try:
        save_auth_request_docx(content, path)
        with open(path, "rb") as f:
            return base64.b64encode(f.read()).decode(), os.path.basename(path)
    finally:
        try:
            os.unlink(path)
        except OSError:
            pass


def draft_auth_request(info: dict[str, Any]) -> str:
    project_name = info.get("项目名称") or info.get("project_name") or "项目"
    return f"关于办理{project_name}授权委托书的请示"


def draft_auth_letter(info: dict[str, Any]) -> str:
    project_name = info.get("项目名称") or info.get("project_name") or "项目"
    return f"{project_name}授权委托书"


def draft_auth_documents(info: dict[str, Any]) -> dict[str, str]:
    with ThreadPoolExecutor(max_workers=2) as executor:
        request_future = executor.submit(draft_auth_request, info)
        letter_future = executor.submit(draft_auth_letter, info)
        return {
            "auth_content": request_future.result(),
            "letter_content": letter_future.result(),
        }


def _create_auth_ledger_workbook():
    import openpyxl
    from openpyxl.styles import Alignment, Border, Font, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "授权委托台账"
    thin = Side(style="thin", color="000000")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for col, header in enumerate(AUTH_LEDGER_HEADERS, start=1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(name="宋体", size=10, bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border
    widths = [8, 18, 12, 12, 12, 24, 8, 42, 46, 12, 18, 18, 12, 36, 12, 8]
    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(idx)].width = width
    ws.row_dimensions[1].height = 57.6
    return wb


def _ensure_ledger_headers(ws) -> None:
    current = [ws.cell(1, col).value for col in range(1, len(AUTH_LEDGER_HEADERS) + 1)]
    if current == AUTH_LEDGER_HEADERS:
        return
    for col, header in enumerate(AUTH_LEDGER_HEADERS, start=1):
        ws.cell(1, col, header)


def init_auth_ledger(ledger_path: str) -> None:
    os.makedirs(os.path.dirname(ledger_path), exist_ok=True)
    with file_lock(ledger_path):
        if os.path.exists(ledger_path):
            return
        atomic_save_workbook(_create_auth_ledger_workbook(), ledger_path)


def build_ledger_row(extracted: dict[str, Any], user_inputs: dict[str, Any], title: str) -> list[Any]:
    a2 = extracted.get("attachment2") or {}
    return [
        None,
        a2.get("authorization_no"),
        user_inputs.get("handler"),
        a2.get("legal_representative"),
        a2.get("trustee_name"),
        user_inputs.get("seal"),
        user_inputs.get("copies"),
        _normalize_authorization_term(a2.get("authorization_term")),
        a2.get("authorization_scope"),
        None,
        None,
        None,
        None,
        title,
        None,
        None,
    ]


def record_to_ledger(extracted: dict[str, Any], user_inputs: dict[str, Any], title: str, ledger_path: str) -> bool:
    import openpyxl
    from openpyxl.styles import Alignment, Border, Font, Side

    os.makedirs(os.path.dirname(ledger_path), exist_ok=True)
    with file_lock(ledger_path):
        if os.path.exists(ledger_path):
            wb = openpyxl.load_workbook(ledger_path)
            ws = wb.active
            _ensure_ledger_headers(ws)
        else:
            wb = _create_auth_ledger_workbook()
            ws = wb.active

        max_seq = 0
        for row in ws.iter_rows(min_row=2, values_only=True):
            try:
                max_seq = max(max_seq, int(row[0] or 0))
            except (TypeError, ValueError):
                pass

        values = build_ledger_row(extracted, user_inputs, title)
        auth_no = values[1]
        target_row = None
        if auth_no:
            normalized_auth_no = re.sub(r"\s+", "", str(auth_no))
            for row_idx in range(2, ws.max_row + 1):
                existing = ws.cell(row_idx, 2).value
                if existing and re.sub(r"\s+", "", str(existing)) == normalized_auth_no:
                    target_row = row_idx
                    values[0] = ws.cell(row_idx, 1).value or max_seq + 1
                    break
        if target_row is None:
            values[0] = max_seq + 1
            target_row = ws.max_row + 1
        thin = Side(style="thin", color="000000")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)
        for col, value in enumerate(values, start=1):
            cell = ws.cell(target_row, col, value)
            cell.font = Font(name="宋体", size=10)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            if col == 9:
                cell.alignment = Alignment(horizontal="justify", vertical="center", wrap_text=True)
            cell.border = border
        ws.row_dimensions[target_row].height = 123
        atomic_save_workbook(wb, ledger_path)
    return True


def import_auth_ledger_rows(rows: list[dict[str, Any]], ledger_path: str) -> dict[str, int]:
    import openpyxl
    from openpyxl.styles import Alignment, Border, Font, Side

    os.makedirs(os.path.dirname(ledger_path), exist_ok=True)
    inserts = 0
    updates = 0
    with file_lock(ledger_path):
        if os.path.exists(ledger_path):
            wb = openpyxl.load_workbook(ledger_path)
            ws = wb.active
            _ensure_ledger_headers(ws)
        else:
            wb = _create_auth_ledger_workbook()
            ws = wb.active

        max_seq = 0
        existing_by_auth_no: dict[str, int] = {}
        for row_idx in range(2, ws.max_row + 1):
            try:
                max_seq = max(max_seq, int(ws.cell(row_idx, 1).value or 0))
            except (TypeError, ValueError):
                pass
            auth_no = ws.cell(row_idx, 2).value
            if auth_no:
                existing_by_auth_no[re.sub(r"\s+", "", str(auth_no))] = row_idx

        thin = Side(style="thin", color="000000")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)
        for row in rows:
            values = [row.get(header) for header in AUTH_LEDGER_HEADERS]
            auth_no = values[1]
            target_row = None
            if auth_no:
                target_row = existing_by_auth_no.get(re.sub(r"\s+", "", str(auth_no)))
            if target_row:
                updates += 1
                values[0] = ws.cell(target_row, 1).value or values[0] or max_seq + 1
            else:
                inserts += 1
                max_seq += 1
                values[0] = values[0] or max_seq
                target_row = ws.max_row + 1
                if auth_no:
                    existing_by_auth_no[re.sub(r"\s+", "", str(auth_no))] = target_row

            for col, value in enumerate(values, start=1):
                cell = ws.cell(target_row, col, value)
                cell.font = Font(name="宋体", size=10)
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                if col == 9:
                    cell.alignment = Alignment(horizontal="justify", vertical="center", wrap_text=True)
                cell.border = border
            ws.row_dimensions[target_row].height = 123

        atomic_save_workbook(wb, ledger_path)
    return {"inserts": inserts, "updates": updates, "count": inserts + updates}


def validate_user_inputs(user_inputs: dict[str, Any]) -> None:
    auth_mode = user_inputs.get("auth_mode")
    if auth_mode not in {"direct", "transfer"}:
        raise ValueError("请选择授权方式")
    if auth_mode == "transfer" and not (user_inputs.get("transfer_subject") or "").strip():
        raise ValueError("转授权需填写转授权委托主体")
    for key, label in [("copies", "授权办理份数"), ("seal", "办理授权使用的章"), ("handler", "经办人")]:
        if not (user_inputs.get(key) or "").strip():
            raise ValueError(f"请填写{label}")
